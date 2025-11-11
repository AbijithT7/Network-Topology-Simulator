import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import networkx as nx
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
import matplotlib.image as mpimg
import numpy as np
import math
from matplotlib.patches import FancyArrow, Rectangle
import webbrowser
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageTk
from PIL import Image, ImageTk


import os

import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))
print("Current working directory:", os.getcwd())

img_pil = Image.open("professor.png")
print("Original size:", img_pil.size)
img_pil = img_pil.resize((120, 120))



# --- Tree layout (robust) ---
def hierarchy_pos(G, root, width=2., vert_gap=0.4, vert_loc=0, xcenter=0.5, pos=None, parent=None):
    if pos is None:
        pos = {root: (xcenter, vert_loc)}
    else:
        pos[root] = (xcenter, vert_loc)
    children = list(G.neighbors(root))
    if parent is not None and parent in children:
        children.remove(parent)
    if len(children) != 0:
        dx = width / len(children)
        nextx = xcenter - width / 2 - dx / 2
        for child in children:
            nextx += dx
            pos = hierarchy_pos(G, child, width=dx, vert_gap=vert_gap,
                                vert_loc=vert_loc - vert_gap, xcenter=nextx, pos=pos, parent=root)
    return pos


# Tooltip implementation
class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tipwindow = None
        widget.bind("<Enter>", self.show)
        widget.bind("<Leave>", self.hide)
    def show(self, event=None):
        if self.tipwindow or not self.text:
            return
        x = self.widget.winfo_rootx() + 25
        y = self.widget.winfo_rooty() + 20
        self.tipwindow = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify=tk.LEFT,
                         background="#ffffe0", relief=tk.SOLID, borderwidth=1,
                         font=("tahoma", 9, "normal"))
        label.pack(ipadx=1)
    def hide(self, event=None):
        tw = self.tipwindow
        self.tipwindow = None
        if tw:
            tw.destroy()


def validate_int(P):
    if P == "": return True
    try: return int(P) > 0
    except ValueError: return False


def validate_float(P):
    if P == "": return True
    try: return float(P) >= 0
    except ValueError: return False


# Globals to save last run inputs/results
saved_graph = None
saved_topology = None
saved_ring = None
saved_result_text = ""
saved_positions = None


# Helper: draw double offset edges with opposite arrows (used for doubly-unidirectional ring)
def draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color="#444"):
    dx = x2 - x1
    dy = y2 - y1
    L = math.hypot(dx, dy)
    if L == 0:
        return
    nxp = -dy / L
    nyp = dx / L
    offset = 0.06
    ax.plot([x1 + nxp*offset, x2 + nxp*offset], [y1 + nyp*offset, y2 + nyp*offset], color=color, linewidth=2, zorder=1)
    ax.plot([x1 - nxp*offset, x2 - nxp*offset], [y1 - nyp*offset, y2 - nyp*offset], color=color, linewidth=2, zorder=1)
    mid1x = (x1 + x2)/2 + nxp*offset
    mid1y = (y1 + y2)/2 + nyp*offset
    arr1 = FancyArrow(mid1x, mid1y, dx*0.0001, dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color='black', length_includes_head=True, zorder=3)
    ax.add_patch(arr1)
    mid2x = (x1 + x2)/2 - nxp*offset
    mid2y = (y1 + y2)/2 - nyp*offset
    arr2 = FancyArrow(mid2x, mid2y, -dx*0.0001, -dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color='black', length_includes_head=True, zorder=3)
    ax.add_patch(arr2)


def build_step_by_step(n, num_edges, port_cost, cable_len, cable_cost):
    lines = []
    lines.append("Step-by-step cost calculation:")
    lines.append(f"1) Number of nodes = {n}")
    lines.append(f"2) Number of connections (edges) = {num_edges}")
    lines.append(f"3) Cost per port = {port_cost}")
    lines.append(f"4) Total port cost = connections * 2 * cost_per_port")
    total_port_cost = num_edges * 2 * port_cost
    lines.append(f"   = {num_edges} * 2 * {port_cost} = {total_port_cost:.2f}")
    lines.append(f"5) Cable length per connection = {cable_len} m")
    lines.append(f"6) Cost per unit cable = {cable_cost} per m")
    total_cable_cost = num_edges * cable_len * cable_cost
    lines.append(f"7) Total cable cost = connections * cable_len * cable_cost")
    lines.append(f"   = {num_edges} * {cable_len} * {cable_cost} = {total_cable_cost:.2f}")
    total_cost = total_port_cost + total_cable_cost
    lines.append(f"8) Total cost = total_port_cost + total_cable_cost = {total_port_cost:.2f} + {total_cable_cost:.2f} = {total_cost:.2f}")
    return "\n".join(lines), total_port_cost, total_cable_cost, total_cost


def generate_topology():
    global saved_graph, saved_topology, saved_ring, saved_result_text, saved_positions
    try:
        n = int(entry_nodes.get())
    except ValueError:
        messagebox.showerror("Input Error", "Enter a positive integer for number of nodes.")
        return


    topology_type = topology_choice.get()
    ring_variant_val = ring_variant.get() if topology_type == "Ring" else None


    if topology_type == "Ring" and ring_variant_val in ["Singly (Unidirectional)", "Doubly (Unidirectional)"]:
        G = nx.DiGraph()
    else:
        G = nx.Graph()
    G.add_nodes_from(range(1, n + 1))


    if topology_type == "Bus":
        for i in range(1, n):
            G.add_edge(i, i + 1)
    elif topology_type == "Star":
        for i in range(2, n + 1):
            G.add_edge(1, i)
    elif topology_type == "Ring":
        for i in range(1, n):
            G.add_edge(i, i + 1)
        G.add_edge(n, 1)
        if ring_variant_val == "Doubly (Unidirectional)":
            for i in range(1, n):
                G.add_edge(i + 1, i)
            G.add_edge(1, n)
    elif topology_type == "Mesh":
        for i in range(1, n + 1):
            for j in range(i + 1, n + 1):
                G.add_edge(i, j)
    elif topology_type == "Tree":
        k = math.log2(n + 1)
        if not k.is_integer():
            messagebox.showerror("Tree Requirement", "Number of nodes for Tree must be 2^k - 1 (e.g., 3,7,15...) for a perfect binary tree.")
            return
        for i in range(2, n + 1):
            G.add_edge(i // 2, i)


    try:
        port_cost = float(entry_port_cost.get())
        cable_len = float(entry_cable_length.get())
        cable_cost = float(entry_cable_cost.get())
    except ValueError:
        messagebox.showerror("Input Error", "Please check cost values and lengths.")
        return


    num_edges = G.number_of_edges()
    step_text, total_port_cost, total_cable_cost, total_cost = build_step_by_step(n, num_edges, port_cost, cable_len, cable_cost)


    summary = (
        f"Topology: {topology_type} {f'({ring_variant_val})' if ring_variant_val else ''}\n"
        f"Nodes: {n}\n"
        f"Connections: {num_edges}\n"
        f"Total Cost: ₹{total_cost:.2f}\n\n"
        f"{step_text}\n"
    )


    result_label.config(state='normal')
    result_label.delete("1.0", tk.END)
    result_label.insert(tk.END, summary)
    result_label.config(state='disabled')


    saved_graph = G
    saved_topology = topology_type
    saved_ring = ring_variant_val
    saved_result_text = summary


    show_graph(G, topology_type, ring_variant_val)


import matplotlib.patheffects as pe  # <-- Add this at the top with other imports

def show_graph(G, topology_type, ring_variant_val, preview=False):
    global saved_positions
    for widget in graph_frame.winfo_children():
        widget.destroy()
    fig, ax = plt.subplots(figsize=(9, 7))

    # --- Layout setup ---
    if topology_type == "Bus":
        nodes = list(G.nodes)
        n = len(nodes)
        xs = np.arange(n)
        ys = np.array([0.6 if i % 2 == 0 else -0.6 for i in range(n)])
        pos = {node: (float(x), float(y)) for node, x, y in zip(nodes, xs, ys)}
        left, right = -0.5, n - 0.5
        ax.plot([left, right], [0, 0], color='black', linewidth=4, zorder=0)
        # Terminators at both ends
        for tx in [left, right]:
            term_rect = Rectangle((tx - 0.09, -0.11), 0.18, 0.22, linewidth=0, edgecolor=None, facecolor='black', zorder=2)
            ax.add_patch(term_rect)
        for x, y in zip(xs, ys):
            ax.plot([x, x], [y - 0.08 if y > 0 else y + 0.08, 0], color='black', linewidth=2, zorder=1)
            ax.plot([x - 0.06, x + 0.06], [0, 0], color='black', linewidth=2, zorder=1)
    elif topology_type == "Ring":
        n = len(G.nodes)
        theta = np.linspace(0, 2 * np.pi, n, endpoint=False)
        pos = {node: (np.cos(th), np.sin(th)) for node, th in zip(G.nodes, theta)}
    elif topology_type == "Tree":
        try:
            pos = hierarchy_pos(G, 1)
        except Exception:
            pos = nx.spring_layout(G, seed=42)
    else:
        pos = nx.spring_layout(G, seed=42)

    saved_positions = pos.copy()

    # --- Draw Edges ---
    if topology_type != "Bus":
        edges_drawn = set()
        for u, v in G.edges():
            x1, y1 = pos[u]
            x2, y2 = pos[v]
            if topology_type == "Ring" and ring_variant_val == "Doubly (Unidirectional)":
                pair = tuple(sorted((u, v)))
                if pair not in edges_drawn:
                    draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color="#444")
                    edges_drawn.add(pair)
            else:
                ax.plot([x1, x2], [y1, y2], color="#444", linewidth=2, zorder=1)
                if topology_type == "Ring" and ring_variant_val == "Singly (Unidirectional)":
                    dx = x2 - x1
                    dy = y2 - y1
                    midpoint = (x1 + 0.5 * dx, y1 + 0.5 * dy)
                    arr = FancyArrow(midpoint[0], midpoint[1], dx * 0.0001, dy * 0.0001,
                                     width=0.008, head_width=0.05, head_length=0.05,
                                     color='black', length_includes_head=True, zorder=3)
                    ax.add_patch(arr)

    # --- Draw Nodes as Computer Icons ---
    try:
        img = mpimg.imread("computer.png")
        node_count = len(pos)
        zoom = 0.12 if node_count <= 8 else (0.09 if node_count <= 16 else 0.06)
        for node, (x, y) in pos.items():
            ab = AnnotationBbox(OffsetImage(img, zoom=zoom), (x, y), frameon=False, zorder=4)
            ax.add_artist(ab)
            # Label number in center of computer
            ax.text(x, y + 0.02, str(node),
                    fontsize=12, fontweight='bold',
                    ha='center', va='center',
                    color='white',
                    path_effects=[pe.withStroke(linewidth=2, foreground='black')],
                    zorder=6)
    except Exception as e:
        nx.draw_networkx_nodes(G, pos, ax=ax, node_size=1100, node_color="#b7e2f9")
        nx.draw_networkx_labels(G, pos, ax=ax, font_size=13, font_weight='bold')

    # --- Graph appearance ---
    ax.set_axis_off()
    plt.title(f"{topology_type}" + (f" ({ring_variant_val})" if ring_variant_val else ""), fontsize=15, weight='bold')

    xs = [p[0] for p in pos.values()]
    ys = [p[1] for p in pos.values()]
    if xs and ys:
        xmin, xmax = min(xs) - 0.9, max(xs) + 0.9
        ymin, ymax = min(ys) - 0.9, max(ys) + 0.9
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(ymin, ymax)

    # --- Embed Matplotlib figure into Tkinter ---
    canvas = FigureCanvasTkAgg(fig, master=graph_frame)
    canvas.draw()
    canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    toolbar = NavigationToolbar2Tk(canvas, graph_frame)
    toolbar.update()
    toolbar.pack(side=tk.BOTTOM, fill=tk.X)
    plt.close(fig)



# Developed By: with prof and black separator
def show_developed_by():
    win = tk.Toplevel(root)
    win.title("Developed By")
    win.config(bg='#f5f7fa')

    tk.Label(win, text="Developed By", font=("Segoe UI", 14, "bold"),
             bg='#f5f7fa').pack(pady=(12, 4), anchor="w", padx=24)

    # Row 1 – Abijith
    row1 = tk.Frame(win, bg='#f5f7fa')
    row1.pack(fill=tk.X, padx=36, pady=(2, 3))
    try:
        abijith_img = Image.open("abijith.jpeg").resize((80, 80))
        abijith_tk = ImageTk.PhotoImage(abijith_img)
        tk.Label(row1, image=abijith_tk, bg='#f5f7fa').pack(side=tk.LEFT, padx=6)
        row1.img_ref = abijith_tk
    except Exception as e:
        tk.Label(row1, text="[Image not found]", bg='#f5f7fa').pack(side=tk.LEFT)
    tk.Label(row1, text="Abijith Thennarasu", font=("Segoe UI", 11),
             bg='#f5f7fa').pack(side=tk.LEFT, padx=12)

    # Row 2 – Dharmayu
    row2 = tk.Frame(win, bg='#f5f7fa')
    row2.pack(fill=tk.X, padx=36, pady=(2, 8))
    try:
        dharmyu_img = Image.open("dharmyu.jpeg").resize((80, 80))
        dharmyu_tk = ImageTk.PhotoImage(dharmyu_img)
        tk.Label(row2, image=dharmyu_tk, bg='#f5f7fa').pack(side=tk.LEFT, padx=6)
        row2.img_ref = dharmyu_tk
    except Exception as e:
        tk.Label(row2, text="[Image not found]", bg='#f5f7fa').pack(side=tk.LEFT)
    tk.Label(row2, text="Dharmayu Jadwani", font=("Segoe UI", 11),
             bg='#f5f7fa').pack(side=tk.LEFT, padx=12)

    # Separator
    sep = tk.Frame(win, height=2, bd=0, relief=tk.SUNKEN, bg="black")
    sep.pack(pady=(5, 15), fill=tk.X, padx=24)

    # Professor section
    row = tk.Frame(win, bg='#f5f7fa')
    row.pack(padx=24, pady=10)
    try:
        img_pil = Image.open("professor.png").resize((120, 120))
        img_tk = ImageTk.PhotoImage(img_pil)
        tk.Label(row, image=img_tk, bg='#f5f7fa').pack(side=tk.LEFT, padx=2)
        row.img_ref = img_tk
    except Exception as e:
        tk.Label(row, text="[image not found]", bg='#f5f7fa').pack(side=tk.LEFT, padx=2)
        print(f"Error loading professor image: {e}")

    prof_desc = tk.Label(row,
                         text="Guided By:\nDr. Swaminathan Annadurai",
                         bg='#f5f7fa',
                         font=("Segoe UI", 12, "bold"),
                         anchor="w", justify='left')
    prof_desc.pack(side=tk.LEFT, padx=(20, 0))
    tk.Button(win, text="Close", command=win.destroy).pack(pady=14)



def show_learn():
    win = tk.Toplevel(root)
    win.title("Learn - Network Topology Concepts")
    win.geometry("950x700")
    win.config(bg='#f5f7fa')

    tk.Label(win, text="Network Topology Concepts", font=("Segoe UI", 16, "bold"),
             bg='#f5f7fa', fg='#007ACC').pack(pady=(12, 8))

    frame = tk.Frame(win, bg='#f5f7fa')
    frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

    scrollbar = tk.Scrollbar(frame)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    text_widget = tk.Text(frame, wrap=tk.WORD, font=("Segoe UI", 10),
                          bg='#ffffff', yscrollcommand=scrollbar.set,
                          padx=15, pady=10)
    text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar.config(command=text_widget.yview)


    learn_content = """
NETWORK TOPOLOGY

═══════════════════════════════════════════════════════════════════════

1. INTRODUCTION TO NETWORK TOPOLOGY

Network topology refers to the arrangement of different elements (links, nodes, etc.) in a computer network. It defines how devices are connected and how data flows between them.

Types of topology:
• Physical Topology: Actual layout of cables and devices
• Logical Topology: How data actually flows in the network

═══════════════════════════════════════════════════════════════════════

2. BUS TOPOLOGY

Structure:
• All devices connect to a single central cable (backbone)
• Data travels in both directions along the cable
• Terminators at both ends prevent signal reflection

Advantages:
✓ Easy to install and extend
✓ Cost-effective (requires less cable)
✓ Suitable for small networks
✓ Easy to understand and implement

Disadvantages:
✗ Single point of failure (if backbone fails, entire network fails)
✗ Performance degrades with more devices
✗ Difficult to troubleshoot
✗ Limited cable length and number of stations

Best Use Cases:
• Small networks (10-12 computers)
• Temporary networks
• Budget-constrained setups

═══════════════════════════════════════════════════════════════════════

3. STAR TOPOLOGY

Structure:
• All devices connect to a central hub/switch
• Each device has dedicated point-to-point connection to hub
• Data passes through the central hub

Advantages:
✓ Easy to install and manage
✓ Failure of one node doesn't affect others
✓ Easy to detect faults and troubleshoot
✓ Easy to add/remove devices
✓ Better performance than bus topology

Disadvantages:
✗ Hub/switch is single point of failure
✗ Requires more cable than bus topology
✗ Cost increases with hub/switch quality
✗ Limited by hub specifications

Best Use Cases:
• Small to medium office networks
• Home networks
• Networks requiring easy management

═══════════════════════════════════════════════════════════════════════

4. RING TOPOLOGY

Structure:
• Each device connects to exactly two other devices
• Forms a circular data path
• Data travels in one or both directions

Variants:
a) Singly Bidirectional: Data can flow both ways on single ring
b) Singly Unidirectional: Data flows in one direction only
c) Doubly Unidirectional: Two separate rings with opposite data flows

Advantages:
✓ Equal access for all devices
✓ No collisions (token passing)
✓ Predictable performance
✓ Can span longer distances than bus

Disadvantages:
✗ Failure of one device/cable affects entire network
✗ Difficult to troubleshoot
✗ Adding/removing devices disrupts network
✗ More expensive than bus topology

Best Use Cases:
• Networks requiring equal access
• Token ring networks (legacy)
• High-speed LANs with token passing

═══════════════════════════════════════════════════════════════════════

5. MESH TOPOLOGY

Structure:
• Every device connects to every other device
• Multiple paths between any two devices
• Provides redundancy and fault tolerance

Types:
a) Full Mesh: Every device connected to every other device
   Connections = n(n-1)/2 where n = number of nodes
   
b) Partial Mesh: Some devices have multiple connections, not all

Advantages:
✓ Highly reliable (multiple paths)
✓ No traffic congestion
✓ Robust and fault-tolerant
✓ Data can be transmitted simultaneously
✓ Privacy and security

Disadvantages:
✗ Very expensive (many cables and ports needed)
✗ Complex installation and maintenance
✗ Requires large amount of cabling
✗ Difficult to configure

Best Use Cases:
• Critical systems requiring high reliability
• Military and financial networks
• Backbone networks
• Wireless mesh networks

═══════════════════════════════════════════════════════════════════════

6. TREE TOPOLOGY

Structure:
• Hierarchical structure like a tree
• Combination of star topologies
• Root node connects to secondary nodes, which connect to tertiary nodes
• Also called Hierarchical Topology

Advantages:
✓ Scalable (easy to expand)
✓ Easy to manage and maintain
✓ Error detection is easy
✓ Suitable for large networks
✓ Point-to-point wiring for segments

Disadvantages:
✗ If backbone fails, entire segment fails
✗ More cables required
✗ Expensive
✗ Difficult to configure

Best Use Cases:
• Large organizations
• Wide Area Networks (WANs)
• Networks with hierarchical structure
• Educational institutions

═══════════════════════════════════════════════════════════════════════

7. COST ANALYSIS FACTORS

When selecting a topology, consider these costs:

1. Port/Interface Costs:
   • Network interface cards (NICs)
   • Switch/hub ports
   • Each connection requires 2 ports

2. Cable Costs:
   • Cable length × cost per unit
   • Different cable types (CAT5e, CAT6, fiber)
   • Installation labor

3. Hardware Costs:
   • Central devices (switches, hubs, routers)
   • Quality and features affect price
   • Redundant devices for fault tolerance

4. Maintenance Costs:
   • Troubleshooting complexity
   • Downtime costs
   • Replacement parts

═══════════════════════════════════════════════════════════════════════

8. COMPARISON SUMMARY

Topology    | Cost    | Reliability | Scalability | Complexity
------------|---------|-------------|-------------|------------
Bus         | Low     | Low         | Poor        | Simple
Star        | Medium  | Medium      | Good        | Simple
Ring        | Medium  | Medium      | Medium      | Medium
Mesh        | High    | Very High   | Poor        | Complex
Tree        | High    | Medium      | Excellent   | Medium

═══════════════════════════════════════════════════════════════════════

9. SELECTION CRITERIA

Choose topology based on:

✓ Network size (number of devices)
✓ Budget constraints
✓ Reliability requirements
✓ Scalability needs
✓ Performance requirements
✓ Ease of maintenance
✓ Physical constraints
✓ Security requirements

═══════════════════════════════════════════════════════════════════════

10. PRACTICAL CONSIDERATIONS

Installation Tips:
• Plan cable routes before installation
• Leave room for expansion
• Label all cables and ports
• Document network layout
• Test connections before deployment
• Consider environmental factors
• Plan for redundancy in critical systems

Troubleshooting:
• Bus: Check terminators and backbone
• Star: Check central hub/switch first
• Ring: Locate break in the ring
• Mesh: Test individual connections
• Tree: Check hierarchical levels systematically

═══════════════════════════════════════════════════════════════════════

REFERENCE VIDEO:
Animated Network Topology Explanation
https://www.youtube.com/watch?v=zbqrNg4C98U

For detailed examples and calculations, use the simulator to experiment with different topologies and cost parameters.

═══════════════════════════════════════════════════════════════════════
    
REFERENCE VIDEO:
Below is the embedded offline video tutorial for Network Topologies.
    """
    text_widget.insert('1.0', learn_content)
    text_widget.config(state='disabled')

    # --- YouTube Video Section ---
    tk.Label(win, text="Reference Video (YouTube):", 
             font=("Segoe UI", 12, "bold"), bg="#f5f7fa", fg="#007ACC").pack(pady=(8, 6))

    def open_youtube_video():
        webbrowser.open("https://www.youtube.com/watch?v=zbqrNg4C98U")

    youtube_btn = tk.Button(
        win, text="▶  Watch on YouTube", 
        command=open_youtube_video,
        font=("Segoe UI", 11, "bold"), 
        bg="#FF0000", fg="white",
        padx=15, pady=8,
        relief="flat", cursor="hand2"
    )
    youtube_btn.pack(pady=(5, 15))

    # Close button
    tk.Button(win, text="Close", command=win.destroy,
              font=("Segoe UI", 10), padx=20, pady=5).pack(pady=10)


def show_help():
    win = tk.Toplevel(root)
    win.title("Help")
    tk.Label(win, text="Help - How to use", font=("Segoe UI", 14, "bold")).pack(pady=8)
    help_text = (
        "How to execute/use:\n\n"
        "1. Enter number of nodes (positive integer).\n"
        "2. Select topology and ring variant (for Ring only).\n"
        "3. Enter cost parameters: cost per port, cable length, cost per unit cable.\n"
        "4. Click 'Generate Topology' to compute cost and view graph.\n\n"
        "Interpretation:\n- Connections: number of links.\n- Total Cost: ports + cable costs (step-by-step shown in output).\n- Bus topology: displays backbone with T-connectors and computers above/below.\n- Ring variants: arrows show direction (single or double).\n\n"
        "To save results: use 'Download Report' to export a detailed DOCX file\n"
        "with inputs, outputs, and embedded graph image.\n"
    )
    tk.Label(win, text=help_text, justify="left", anchor="w").pack(padx=12, pady=6)
    tk.Button(win, text="Close", command=win.destroy).pack(pady=6)


def save_graph_to_file(filepath):
    """Helper function to save graph image to file"""
    if saved_graph is None or saved_positions is None:
        return False
    
    G = saved_graph
    topology_type = saved_topology
    ring_variant_val = saved_ring
    pos = saved_positions
    
    fig = plt.figure(figsize=(10, 8))
    ax = fig.add_subplot(111)
    
    if topology_type == "Bus":
        left, right = -0.5, len(pos) - 0.5
        ax.plot([left, right], [0, 0], color='black', linewidth=4, zorder=0)
        for tx in [left, right]:
            term_rect = Rectangle((tx - 0.09, -0.11), 0.18, 0.22, linewidth=0, edgecolor=None, facecolor='black', zorder=2)
            ax.add_patch(term_rect)
        for node, (x, y) in pos.items():
            ax.plot([x, x], [y - 0.08 if y > 0 else y + 0.08, 0], color='black', linewidth=2)
            ax.plot([x - 0.06, x + 0.06], [0, 0], color='black', linewidth=2)
    
    edges_drawn = set()
    if topology_type != "Bus":
        for u, v in G.edges():
            x1, y1 = pos[u]; x2, y2 = pos[v]
            if topology_type == "Ring" and ring_variant_val == "Doubly (Unidirectional)":
                pair = tuple(sorted((u, v)))
                if pair not in edges_drawn:
                    draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color="#444")
                    edges_drawn.add(pair)
            else:
                ax.plot([x1, x2], [y1, y2], color="#444", linewidth=2)
                if topology_type == "Ring" and ring_variant_val == "Singly (Unidirectional)":
                    dx = x2 - x1; dy = y2 - y1
                    mid = (x1 + 0.5*dx, y1 + 0.5*dy)
                    arr = FancyArrow(mid[0], mid[1], dx*0.0001, dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color='black', length_includes_head=True)
                    ax.add_patch(arr)
    
    try:
        img = mpimg.imread("computer.png")
        node_count = len(pos)
        zoom = 0.14 if node_count <= 8 else (0.1 if node_count <= 16 else 0.07)
        for node, (x, y) in pos.items():
            ab = AnnotationBbox(OffsetImage(img, zoom=zoom), (x, y), frameon=False)
            ax.add_artist(ab)
    except Exception:
        nx.draw_networkx_nodes(G, pos, ax=ax, node_size=1100, node_color="#b7e2f9")
    
    if topology_type == "Bus":
        for node, (x, y) in pos.items():
            ypos = y + 0.15 if y > 0 else y - 0.15
            ax.text(x, ypos, str(node), fontsize=12, fontweight='bold', ha='center', va='center')
    else:
        nx.draw_networkx_labels(G, pos, ax=ax, font_size=13, font_weight='bold')
    
    ax.set_axis_off()
    plt.title(f"{topology_type}" + (f" ({ring_variant_val})" if ring_variant_val else ""), fontsize=15, weight='bold')
    
    fig.savefig(filepath, bbox_inches='tight', dpi=150)
    plt.close(fig)
    return True


def do_download():
    """Create comprehensive DOCX report with inputs, outputs, and graph image"""
    if saved_graph is None:
        messagebox.showinfo("Nothing to download", "Generate a topology first, then download.")
        return
    
    try:
        # Ask for save location
        fpath = filedialog.asksaveasfilename(
            defaultextension=".docx", 
            filetypes=[("Word Document","*.docx")], 
            title="Save detailed report as..."
        )
        if not fpath:
            return
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('Network Topology Simulator - Detailed Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.runs[0].font
        title_format.color.rgb = RGBColor(0, 122, 204)
        
        doc.add_paragraph()
        
        # Input Parameters Section
        doc.add_heading('Input Parameters', 1)
        topology_type = saved_topology or ""
        ring_variant_val = saved_ring or ""
        
        input_table = doc.add_table(rows=6, cols=2)
        input_table.style = 'Light Grid Accent 1'
        
        input_data = [
            ('Topology Type', f"{topology_type} {('('+ring_variant_val+')') if ring_variant_val else ''}"),
            ('Number of Nodes', entry_nodes.get()),
            ('Cost per Port (₹)', entry_port_cost.get()),
            ('Cable Length per Connection (m)', entry_cable_length.get()),
            ('Cost per Unit Cable (₹/m)', entry_cable_cost.get()),
            ('Total Connections', str(saved_graph.number_of_edges()))
        ]
        
        for i, (param, value) in enumerate(input_data):
            input_table.rows[i].cells[0].text = param
            input_table.rows[i].cells[1].text = value
            # Bold the parameter names
            input_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Results Section
        doc.add_heading('Cost Analysis Results', 1)
        
        # Add the step-by-step calculation
        result_para = doc.add_paragraph()
        result_para.add_run(saved_result_text).font.name = 'Consolas'
        
        doc.add_paragraph()
        
        # Graph Image Section
        doc.add_heading('Network Topology Diagram', 1)
        
        # Save graph to temporary file
        import tempfile
        temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_img_path = temp_img.name
        temp_img.close()
        
        if save_graph_to_file(temp_img_path):
            try:
                doc.add_picture(temp_img_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Error embedding graph image: {e}]")
            finally:
                # Clean up temp file
                import os
                try:
                    os.unlink(temp_img_path)
                except:
                    pass
        else:
            doc.add_paragraph("[Graph image could not be generated]")
        
        doc.add_paragraph()
        
        # Connection List Section
        doc.add_heading('Network Connections', 1)
        edges_para = doc.add_paragraph()
        
        for u, v in saved_graph.edges():
            edges_para.add_run(f"• Node {u} ↔ Node {v}\n")
        
        doc.add_paragraph()
        
        # Footer
        footer_para = doc.add_paragraph()
        footer_para.add_run('\n' + '─' * 70 + '\n').font.color.rgb = RGBColor(128, 128, 128)
        footer_run = footer_para.add_run('Generated by Network Topology Simulator\n')
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the document
        doc.save(fpath)
        messagebox.showinfo("Success", f"Detailed report with graph saved to:\n{fpath}")
        
    except Exception as e:
        messagebox.showerror("Error", f"Could not create report: {e}\n\nMake sure python-docx is installed:\npip install python-docx")


# ------------------- GUI -------------------
root = tk.Tk()
root.title("Network Topology Simulator & Cost Estimator")
root.geometry("1200x780")
root.minsize(900, 650)
style = ttk.Style(root)
style.theme_use("clam")
root.configure(bg="#f5f7fa")


# Header frame + top-right buttons
frame_header = tk.Frame(root, bg="#007ACC")
frame_header.pack(fill=tk.X)
tk.Label(frame_header, text=" ⛓ Network Topology Simulator ", font=("Segoe UI", 18, "bold"), fg="white", bg="#007ACC", pady=10).pack(side=tk.LEFT, padx=12)
btn_frame = tk.Frame(frame_header, bg="#007ACC")
btn_frame.pack(side=tk.RIGHT, padx=10)
tk.Button(btn_frame, text="Developed By", command=show_developed_by, bg="#f7f7f7").pack(side=tk.LEFT, padx=6, pady=8)
tk.Button(btn_frame, text="Learn", command=show_learn, bg="#f7f7f7").pack(side=tk.LEFT, padx=6, pady=8)
tk.Button(btn_frame, text="Help", command=show_help, bg="#f7f7f7").pack(side=tk.LEFT, padx=6, pady=8)
tk.Button(btn_frame, text="Download Report", command=do_download, bg="#28a745", fg="white", font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT, padx=6, pady=8)


main_frame = tk.Frame(root, bg="#f5f7fa")
main_frame.pack(fill=tk.BOTH, expand=True)
input_panel = tk.Frame(main_frame, width=430, bg="#f5f7fa")
input_panel.pack(side=tk.LEFT, fill=tk.Y, padx=12, pady=10)
header = tk.Label(input_panel, text="🛠 Network Topology Simulator", font=("Segoe UI", 18, "bold"),
                  bg="#f5f7fa", fg="#007ACC")
header.pack(anchor="w", padx=2, pady=(8, 6))


frame_main = tk.LabelFrame(input_panel, text="Design Parameters", font=("Segoe UI", 13, "bold"), bg="#f5f7fa")
frame_main.pack(fill=tk.X, padx=2, pady=6)
vcmd_int = root.register(validate_int)
vcmd_float = root.register(validate_float)
tk.Label(frame_main, text="Number of Nodes:", bg="#f5f7fa", font=("Segoe UI", 11)).grid(row=0, column=0, sticky="w", padx=5, pady=6)
entry_nodes = tk.Entry(frame_main, width=10, validate="key", validatecommand=(vcmd_int, "%P"), font=("Segoe UI", 11))
entry_nodes.grid(row=0, column=1, padx=5, pady=6)
ToolTip(entry_nodes, "Enter a positive integer (e.g., 6, 12)")
tk.Label(frame_main, text="Select Topology:", bg="#f5f7fa", font=("Segoe UI", 11)).grid(row=1, column=0, sticky="w", padx=5, pady=6)
topology_choice = ttk.Combobox(frame_main, values=["Bus", "Star", "Ring", "Mesh", "Tree"], width=17, font=("Segoe UI", 11))
topology_choice.grid(row=1, column=1, padx=5, pady=6)
topology_choice.current(0)
tk.Label(frame_main, text="Ring Variant:", bg="#f5f7fa", font=("Segoe UI", 11)).grid(row=2, column=0, sticky="w", padx=5, pady=6)
ring_variant = ttk.Combobox(frame_main, values=[
        "Singly (Bidirectional)", "Singly (Unidirectional)", "Doubly (Unidirectional)"
    ], width=17, font=("Segoe UI", 11), state="readonly")
ring_variant.grid(row=2, column=1, padx=5, pady=6)
ring_variant.current(0)
def update_ring_dropdown(event=None):
    if topology_choice.get() == "Ring":
        ring_variant.config(state="readonly")
    else:
        ring_variant.config(state="disabled")
topology_choice.bind("<<ComboboxSelected>>", update_ring_dropdown)
update_ring_dropdown()


frame_cost = tk.LabelFrame(input_panel, text="Cost Parameters", font=("Segoe UI", 13, "bold"), bg="#f5f7fa")
frame_cost.pack(fill=tk.X, padx=2, pady=6)
tk.Label(frame_cost, text="Cost per Port (₹):", bg="#f5f7fa", font=("Segoe UI", 11)).grid(row=0, column=0, sticky="w", padx=6, pady=5)
entry_port_cost = tk.Entry(frame_cost, width=10, validate="key", validatecommand=(vcmd_float, "%P"), font=("Segoe UI", 11))
entry_port_cost.grid(row=0, column=1, padx=6, pady=5)
entry_port_cost.insert(0, "100")
tk.Label(frame_cost, text="Cable Length per Connection (m):", bg="#f5f7fa", font=("Segoe UI", 11)).grid(row=1, column=0, sticky="w", padx=6, pady=5)
entry_cable_length = tk.Entry(frame_cost, width=10, validate="key", validatecommand=(vcmd_float, "%P"), font=("Segoe UI", 11))
entry_cable_length.grid(row=1, column=1, padx=6, pady=5)
entry_cable_length.insert(0, "10")
tk.Label(frame_cost, text="Cost per Unit Cable (₹/m):", bg="#f5f7fa", font=("Segoe UI", 11)).grid(row=2, column=0, sticky="w", padx=6, pady=5)
entry_cable_cost = tk.Entry(frame_cost, width=10, validate="key", validatecommand=(vcmd_float, "%P"), font=("Segoe UI", 11))
entry_cable_cost.grid(row=2, column=1, padx=6, pady=5)
entry_cable_cost.insert(0, "50")
tk.Button(input_panel, text="Generate Topology", font=("Segoe UI", 12, "bold"), bg="#27a745", fg="white", command=generate_topology).pack(pady=10)
frame_results = tk.LabelFrame(input_panel, text="Output & Step-by-Step", font=("Segoe UI", 13, "bold"), bg="#f5f7fa")
frame_results.pack(fill=tk.BOTH, expand=False, padx=2, pady=6)
result_label = tk.Text(frame_results, font=("Consolas", 11), bg="#f3f3f3", height=13, width=38, state="disabled", bd=0, wrap="word")
result_label.pack(padx=7, pady=5, fill=tk.BOTH, expand=True)
graph_frame = tk.Frame(main_frame, bg="#f5f7fa")
graph_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=0, pady=0)


root.mainloop()
